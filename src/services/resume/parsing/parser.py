import re
from io import BytesIO
from typing import Dict, Any

from src.domain.ats_models import ExperienceEntry, ProjectEntry, ResumeData

try:
    import fitz
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None
try:
    import docx
except ImportError:
    docx = None

RESPONSE_SCHEMA = {
    "professional_summary": ["..."],
    "skills": ["..."],
    "education": ["..."],
    "experience": ["..."],
    "projects": ["..."],
    "certifications": ["..."],
    "achievements": ["..."],
}

CLEANING_OUTPUT_SCHEMA = {
    "professional_summary_seed": "...",
    "skills": ["..."],
    "education": ["Degree | Institution | Duration | Location | Details"],
    "experience": [
        {
            "role": "...",
            "company": "...",
            "duration": "...",
            "location": "...",
            "highlights": ["..."],
        }
    ],
    "projects": [
        {
            "name": "...",
            "technologies": "...",
            "year": "...",
            "highlights": ["..."],
        }
    ],
    "certifications": ["..."],
    "achievements": ["..."],
    "job_description_keywords": ["..."],
}

def extract_text_from_pdf(file_path: str) -> str:
    if not pdfplumber:
        raise ImportError("pdfplumber is not installed.")
    with pdfplumber.open(file_path) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)

def extract_text_from_docx(file_path: str) -> str:
    if not docx:
        raise ImportError("python-docx is not installed.")
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def parse_resume_text(text: str) -> Dict[str, Any]:
    # Simple regex-based section extraction (can be improved)
    sections = {
        "professional_summary": [],
        "skills": [],
        "education": [],
        "experience": [],
        "projects": [],
        "certifications": [],
        "achievements": [],
    }
    current_section = None
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        lline = line.lower()
        if "summary" in lline:
            current_section = "professional_summary"
        elif "skill" in lline:
            current_section = "skills"
        elif "education" in lline:
            current_section = "education"
        elif "experience" in lline:
            current_section = "experience"
        elif "project" in lline:
            current_section = "projects"
        elif "certification" in lline:
            current_section = "certifications"
        elif "achievement" in lline:
            current_section = "achievements"
        elif current_section:
            sections[current_section].append(line)
    return sections

def parse_resume_file(file_path: str) -> Dict[str, Any]:
    if file_path.lower().endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type")
    return parse_resume_text(text)


def parse_resume(file_bytes: bytes, mime_type: str) -> ResumeData:
    text = _extract_text_from_bytes(file_bytes=file_bytes, mime_type=mime_type)
    section_map = _extract_section_map(text)
    skills = _extract_skills(section_map.get("skills", ""))
    if len(skills) < 3:
        skills = _merge_skill_lists(skills, _infer_skills_from_text(text))
    experience = _extract_experience(section_map.get("experience", ""))
    if not experience:
        experience = _extract_experience(text)
    projects = _extract_projects(section_map.get("projects", ""))
    if not projects:
        projects = _extract_projects(text)
    education = [line.strip() for line in section_map.get("education", "").splitlines() if line.strip()]

    return ResumeData(
        skills=skills,
        experience=experience,
        projects=projects,
        education=education,
        raw_text=text,
        section_map=section_map,
    )


def _extract_text_from_bytes(file_bytes: bytes, mime_type: str) -> str:
    mime = (mime_type or "").lower()
    if "pdf" in mime:
        if not fitz:
            raise RuntimeError("PyMuPDF is not installed.")
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return "\n".join(page.get_text("text") for page in doc)

    if "word" in mime or "docx" in mime:
        if not docx:
            raise RuntimeError("python-docx is not installed.")
        document = docx.Document(BytesIO(file_bytes))
        return "\n".join(para.text for para in document.paragraphs)

    raise ValueError("Unsupported mime type. Only PDF and DOCX are allowed.")


def _extract_section_map(text: str) -> dict[str, str]:
    section_aliases = {
        "summary": ["summary", "professional summary", "profile", "objective"],
        "skills": ["skills", "technical skills", "core skills"],
        "experience": ["experience", "work experience", "employment history"],
        "projects": ["projects", "project experience"],
        "education": ["education", "academic background"],
    }
    ordered_keys = ["summary", "skills", "experience", "projects", "education"]

    current = "summary"
    sections: dict[str, list[str]] = {key: [] for key in ordered_keys}

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        lower = line.lower().strip(":")
        normalized_heading = re.sub(r"[^a-z\s]", "", lower.replace("#", " ")).strip()
        is_short_heading = 1 <= len(normalized_heading.split()) <= 4

        switched = False
        for key, aliases in section_aliases.items():
            if is_short_heading and any(normalized_heading == alias for alias in aliases):
                current = key
                switched = True
                break

            # Support inline heading formats like "Skills: Python, SQL".
            for alias in aliases:
                prefix = f"{alias}:"
                if lower.startswith(prefix):
                    current = key
                    remainder = line[len(prefix) :].strip()
                    if remainder:
                        sections[current].append(remainder)
                    switched = True
                    break
            if switched:
                break

        if not switched:
            sections[current].append(line)

    return {
        "summary": "\n".join(sections["summary"]).strip(),
        "skills": "\n".join(sections["skills"]).strip(),
        "experience": "\n".join(sections["experience"]).strip(),
        "projects": "\n".join(sections["projects"]).strip(),
        "education": "\n".join(sections["education"]).strip(),
    }


def _extract_skills(skills_text: str) -> list[str]:
    values = []
    for token in re.split(r"[,|\n]", skills_text):
        item = token.strip()
        if item:
            values.append(item)
    seen = set()
    ordered = []
    for value in values:
        normalized = value.lower()
        if normalized in seen:
            continue
        seen.add(normalized)
        ordered.append(value)
    return ordered


def _infer_skills_from_text(text: str) -> list[str]:
    detected: list[str] = []
    lowered = (text or "").lower()

    common_skills = [
        "python",
        "java",
        "javascript",
        "typescript",
        "react",
        "node.js",
        "fastapi",
        "rest api",
        "sql",
        "postgresql",
        "mysql",
        "mongodb",
        "redis",
        "docker",
        "kubernetes",
        "aws",
        "gcp",
        "azure",
        "pytorch",
        "tensorflow",
        "machine learning",
        "deep learning",
    ]

    for skill in common_skills:
        if skill in lowered:
            detected.append(skill)

    return detected


def _merge_skill_lists(primary: list[str], fallback: list[str]) -> list[str]:
    merged = list(primary)
    seen = {item.lower() for item in merged}
    for item in fallback:
        normalized = item.lower()
        if normalized in seen:
            continue
        seen.add(normalized)
        merged.append(item)
    return merged


def _extract_experience(experience_text: str) -> list[ExperienceEntry]:
    entries: list[ExperienceEntry] = []
    blocks = [block.strip() for block in re.split(r"\n\s*\n", experience_text) if block.strip()]
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        header = lines[0]
        parts = [part.strip() for part in header.split("|")]
        bullets = [line.lstrip("-*• ").strip() for line in lines[1:] if line.strip()]
        entries.append(
            ExperienceEntry(
                title=parts[0] if len(parts) > 0 else "",
                company=parts[1] if len(parts) > 1 else "",
                duration=parts[2] if len(parts) > 2 else "",
                location=parts[3] if len(parts) > 3 else "",
                bullets=bullets,
            )
        )
    return entries


def _extract_projects(project_text: str) -> list[ProjectEntry]:
    entries: list[ProjectEntry] = []
    blocks = [block.strip() for block in re.split(r"\n\s*\n", project_text) if block.strip()]
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        header = lines[0]
        parts = [part.strip() for part in header.split("|")]
        description = "\n".join(line.lstrip("-*• ").strip() for line in lines[1:])
        technologies = [token.strip() for token in re.split(r"[,/]| and ", parts[1] if len(parts) > 1 else "") if token.strip()]
        entries.append(
            ProjectEntry(
                name=parts[0] if len(parts) > 0 else "",
                technologies=technologies,
                description=description,
            )
        )
    return entries
